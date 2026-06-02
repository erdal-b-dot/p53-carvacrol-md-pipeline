"""
Generate Methods section Word document (.docx)
Q1-level academic English — Journal of Chemical Information and Modeling style
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1.0)
section.right_margin  = Inches(1.0)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)

# ── Styles ──────────────────────────────────────────────────────────
normal_style = doc.styles['Normal']
normal_style.font.name = 'Times New Roman'
normal_style.font.size = Pt(12)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    if level == 1:
        run.font.size = Pt(14)
    elif level == 2:
        run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    return p

def add_para(doc, text, bold=False, italic=False, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(22)  # ~1.5 line spacing
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold   = bold
    run.italic = italic
    return p

def add_mixed_para(doc, parts, indent=True):
    """parts: list of (text, bold, italic) tuples"""
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(22)
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.3)
    for text, bold, italic in parts:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold   = bold
        run.italic = italic
    return p

# ════════════════════════════════════════════════════════════════════
# TITLE
# ════════════════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(4)
tr = title.add_run(
    "Molecular Dynamics Simulation of Carvacrol Binding to the p53 "
    "DNA-Binding Domain: Computational Methods"
)
tr.font.name = 'Times New Roman'
tr.font.size = Pt(14)
tr.bold = True

subtitle = doc.add_paragraph()
subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(16)
sr = subtitle.add_run(
    "[Methods Section — draft for submission]"
)
sr.font.name = 'Times New Roman'
sr.font.size = Pt(11)
sr.italic = True
sr.font.color.rgb = RGBColor(100, 100, 100)

# ════════════════════════════════════════════════════════════════════
# 2. MATERIALS AND METHODS
# ════════════════════════════════════════════════════════════════════
add_heading(doc, "2. Materials and Methods", level=1)

# ── 2.1 ─────────────────────────────────────────────────────────────
add_heading(doc, "2.1. Protein Structure Preparation", level=2)

add_mixed_para(doc, [
    ("The crystal structure of the p53 tumor suppressor protein core domain "
     "in complex with DNA was retrieved from the RCSB Protein Data Bank "
     "(PDB ID: ", False, False),
    ("1TSR", False, True),
    ("; resolution: 2.35 Å; Cho et al., 1994). The asymmetric unit contains "
     "two p53 core domain chains (chains A and B) bound to a double-stranded "
     "DNA duplex. Prior to simulation, all DNA chains, crystallographic water "
     "molecules, and non-structural heteroatoms were removed from the "
     "coordinate file using an automated Python script. The structural zinc "
     "ion (Zn", False, False),
    ("2+", False, False),
    ("), which is coordinated by residues Cys176, His179, Cys238, and Cys242 "
     "and is essential for maintaining the structural integrity of the "
     "DNA-binding loop, was explicitly retained in the final protein model.",
     False, False),
])

add_mixed_para(doc, [
    ("Protonation states of all titratable residues were determined at "
     "physiological pH (7.4) using ", False, False),
    ("PROPKA3", False, True),
    (" (v3.5.1; Olsson et al., 2011). Particular attention was given to "
     "histidine residues, whose protonation state (HID: N", False, False),
    ("δ", False, False),
    ("-protonated; HIE: N", False, False),
    ("ε", False, False),
    ("-protonated; HIP: doubly protonated) can substantially affect the "
     "electrostatic environment of the active site. Residues with predicted "
     "pK", False, False),
    ("a", False, True),
    (" values deviating by more than 1 pH unit from the standard value were "
     "assigned non-default protonation states. The zinc-coordinating "
     "histidine (His179) was assigned the HIE state, consistent with its "
     "role as a metal ligand.", False, False),
])

# ── 2.2 ─────────────────────────────────────────────────────────────
add_heading(doc, "2.2. Ligand Preparation and Force Field Parametrization", level=2)

add_mixed_para(doc, [
    ("The three-dimensional structure of carvacrol (2-methyl-5-(propan-2-yl)"
     "phenol; systematic name: 5-isopropyl-2-methylphenol; PubChem CID: "
     "10364; molecular formula: C", False, False),
    ("10", False, False),
    ("H", False, False),
    ("14", False, False),
    ("O; MW: 150.22 g/mol; SMILES: ", False, False),
    ("Cc1ccc(C(C)C)cc1O", True, False),
    (") was obtained from the PubChem Compound Database in SDF format "
     "(record type: 3D). Force field parameters for carvacrol were "
     "generated using ", False, False),
    ("ACPYPE", False, True),
    (" (AnteChamber PYthon Parser interfacE; Sousa da Silva and Vranken, "
     "2012) with the General AMBER Force Field version 2 (GAFF2; Wang et al., "
     "2004) atom-type assignment. Partial atomic charges were derived using "
     "the AM1-BCC semiempirical quantum mechanical charge model "
     "(Jakalian et al., 2002), which provides a computationally efficient "
     "approximation of HF/6-31G* RESP charges. The net formal charge of "
     "carvacrol was set to zero (neutral species at pH 7.4, well above its "
     "phenolic pK", False, False),
    ("a", False, True),
    (" of ~10.3). The residue name was standardized to ", False, False),
    ("LIG", True, False),
    (" for compatibility with the GROMACS topology framework.", False, False),
])

# ── 2.3 ─────────────────────────────────────────────────────────────
add_heading(doc, "2.3. Molecular Docking", level=2)

add_mixed_para(doc, [
    ("The initial binding pose of carvacrol within the p53 core domain was "
     "determined by molecular docking using ", False, False),
    ("AutoDock Vina", False, True),
    (" 1.2.6 (Eberhardt et al., 2021). Receptor and ligand input files in "
     "PDBQT format were prepared using ", False, False),
    ("Meeko", False, True),
    (" 0.7.1 (Forli laboratory, The Scripps Research Institute). To ensure "
     "unbiased binding site identification, a blind docking protocol was "
     "applied in which the search space encompassed the entire solvent-exposed "
     "surface of the p53 core domain. The docking grid box was centered on "
     "the center of mass of the protein and extended 10 Å beyond the "
     "maximum atomic coordinates in each Cartesian direction. The "
     "exhaustiveness parameter was set to 32 to ensure thorough conformational "
     "sampling, and ten binding modes were requested. The random seed was "
     "fixed (seed = 42) to ensure reproducibility. The top-ranked docking "
     "pose (lowest predicted binding free energy) was selected as the "
     "starting conformation for molecular dynamics simulations.", False, False),
])

# ── 2.4 ─────────────────────────────────────────────────────────────
add_heading(doc, "2.4. Molecular Dynamics System Setup", level=2)

add_mixed_para(doc, [
    ("All molecular dynamics simulations were performed using ", False, False),
    ("GROMACS", False, True),
    (" 2024.1 (Abraham et al., 2015), compiled with Intel oneAPI 2024 "
     "optimizations. The protein was described by the AMBER14SB force field "
     "(Maier et al., 2015), and the ligand was described by GAFF2 parameters "
     "as described in Section 2.2. The protein–ligand complex was assembled "
     "by combining the GROMACS coordinate files generated by ", False, False),
    ("pdb2gmx", True, False),
    (" and ACPYPE, respectively.", False, False),
])

add_mixed_para(doc, [
    ("The system was solvated in a rhombic dodecahedral periodic box with a "
     "minimum distance of 1.2 nm between any solute atom and the box boundary. "
     "The TIP3P water model (Jorgensen et al., 1983) was employed. The "
     "overall charge of the system was neutralized by replacing water "
     "molecules with Na", False, False),
    ("+", False, False),
    (" or Cl", False, False),
    ("−", False, False),
    (" ions as appropriate, and the physiological ionic strength of 0.15 M "
     "NaCl was established using the GROMACS ", False, False),
    ("genion", True, False),
    (" utility.", False, False),
])

# ── 2.5 ─────────────────────────────────────────────────────────────
add_heading(doc, "2.5. Energy Minimization and Equilibration", level=2)

add_mixed_para(doc, [
    ("Prior to production dynamics, the solvated system was subjected to "
     "energy minimization using the steepest descent algorithm until the "
     "maximum force on any atom fell below 10 kJ mol", False, False),
    ("−1", False, False),
    (" nm", False, False),
    ("−1", False, False),
    (" (maximum 50,000 steps). Subsequently, the system was equilibrated in "
     "two sequential phases. In the first phase, a 100 ps constant-volume "
     "(NVT) simulation was performed at 310 K using the V-rescale velocity-"
     "rescaling thermostat (Bussi et al., 2007; coupling time constant "
     "τ", False, False),
    ("T", False, True),
    (" = 0.1 ps) with harmonic position restraints applied to all heavy "
     "atoms of the protein and ligand (force constant k = 1000 kJ mol", False, False),
    ("−1", False, False),
    (" nm", False, False),
    ("−2", False, False),
    ("). Initial velocities were assigned from a Maxwell–Boltzmann "
     "distribution at 310 K.", False, False),
])

add_mixed_para(doc, [
    ("In the second equilibration phase, an isothermal–isobaric (NPT) "
     "simulation of 1 ns was conducted at 310 K and 1.0 bar using the same "
     "thermostat settings combined with the Parrinello–Rahman barostat "
     "(Parrinello and Rahman, 1981; coupling time constant τ", False, False),
    ("P", False, True),
    (" = 2.0 ps; isothermal compressibility = 4.5 × 10", False, False),
    ("−5", False, False),
    (" bar", False, False),
    ("−1", False, False),
    ("). Position restraints were maintained throughout both equilibration "
     "phases to allow solvent and ions to relax around the fixed solute. "
     "The protein and ligand were coupled to one thermostat group "
     "(", False, False),
    ("Protein_LIG", True, False),
    ("), and the solvent was coupled to a separate thermostat group "
     "(", False, False),
    ("Water_and_ions", True, False),
    ("), in accordance with recommended best practices for heterogeneous "
     "biomolecular systems.", False, False),
])

# ── 2.6 ─────────────────────────────────────────────────────────────
add_heading(doc, "2.6. Production Molecular Dynamics Simulation", level=2)

add_mixed_para(doc, [
    ("Following equilibration, unrestrained production MD was performed "
     "for 1 µs (1,000,000 ps; 500,000,000 integration steps) using a "
     "2 fs timestep. The simulation was carried out on the Turkish National "
     "High-Performance Computing infrastructure (TRUBA; https://www.truba.gov.tr) "
     "using the GROMACS 2024.1-oneapi-2024 build. The SLURM job scheduler "
     "was used to allocate resources on the ", False, False),
    ("hamsi", True, False),
    (" partition (AMD EPYC 7742 processors, 128 cores per node). A single "
     "MPI rank with 56 OpenMP threads was employed, which was found to be "
     "optimal for the system size (~50,000 atoms). The total simulation "
     "was divided into consecutive SLURM jobs (72-hour wall-clock limit per "
     "job), each automatically resuming from a binary checkpoint file "
     "(", False, False),
    ("md.cpt", True, False),
    (") to ensure seamless continuation.", False, False),
])

# ── 2.7 ─────────────────────────────────────────────────────────────
add_heading(doc, "2.7. Simulation Parameters", level=2)

add_mixed_para(doc, [
    ("Throughout all MD phases, electrostatic interactions were calculated "
     "using the Particle Mesh Ewald (PME) method (Darden et al., 1993; "
     "Essmann et al., 1995) with a real-space cutoff of 1.2 nm and a "
     "Fourier grid spacing of 0.12 nm. Van der Waals interactions were "
     "treated with a force-switch modifier applied between 1.0 and 1.2 nm, "
     "consistent with AMBER force field recommendations (Shirts et al., 2020). "
     "Long-range dispersion corrections for energy and pressure were applied "
     "using the isotropic ",False, False),
    ("DispCorr", True, False),
    (" = EnerPres setting. All covalent bonds involving hydrogen atoms were "
     "constrained using the LINCS algorithm (Hess et al., 1997; "
     "expansion order 4, two iterations), enabling the use of a 2 fs "
     "timestep. Neighbor lists were updated every 20 steps using the "
     "Verlet cutoff scheme.", False, False),
])

# ── 2.8 ─────────────────────────────────────────────────────────────
add_heading(doc, "2.8. Trajectory Analysis", level=2)

add_mixed_para(doc, [
    ("Trajectory analysis will be performed using standard GROMACS analysis "
     "utilities and custom Python scripts (MDAnalysis v2.x; Michaud-Agrawal "
     "et al., 2011; Gowers et al., 2016). The following structural and "
     "dynamic properties will be computed over the production trajectory: "
     "(i) backbone root-mean-square deviation (RMSD) relative to the "
     "energy-minimized structure, (ii) per-residue root-mean-square "
     "fluctuations (RMSF), (iii) protein–ligand interaction energies "
     "(van der Waals and electrostatic contributions decomposed), "
     "(iv) ligand binding pose stability assessed by heavy-atom RMSD of "
     "carvacrol relative to the docked pose, (v) radius of gyration (R",
     False, False),
    ("g", False, True),
    (") of the protein, and (vi) solvent-accessible surface area (SASA) "
     "of the binding interface. Hydrogen bond formation between carvacrol "
     "and p53 residues will be monitored using a donor–acceptor distance "
     "cutoff of 3.5 Å and an angle cutoff of 30°. Principal component "
     "analysis (PCA) of backbone Cα atoms will be used to characterize "
     "dominant conformational motions.", False, False),
])

# ── 2.9 ─────────────────────────────────────────────────────────────
add_heading(doc, "2.9. Software and Reproducibility", level=2)

add_mixed_para(doc, [
    ("All preparatory scripts, force field parameter files, MDP input "
     "files, and SLURM submission scripts are deposited in a publicly "
     "accessible GitHub repository (URL to be provided upon acceptance). "
     "The pipeline is fully automated and reproducible: executing the "
     "scripts in sequential order (Steps 1–5) yields an identical "
     "simulation-ready system from the raw PDB file. "
     "Software versions used in this study are summarized in Table 1.",
     False, False),
])

doc.add_paragraph()  # spacing

# ── Table 1 ──────────────────────────────────────────────────────────
table_heading = doc.add_paragraph()
table_heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
thr = table_heading.add_run(
    "Table 1. Software packages and versions used in this study."
)
thr.font.name  = 'Times New Roman'
thr.font.size  = Pt(11)
thr.bold = True

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'

header_cells = table.rows[0].cells
for cell, text in zip(header_cells,
                      ["Software", "Version", "Function", "Reference"]):
    cell.text = text
    run = cell.paragraphs[0].runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.bold = True

rows_data = [
    ("GROMACS",       "2024.1-oneapi-2024", "MD engine",                 "Abraham et al., 2015"),
    ("AMBER14SB FF",  "—",                  "Protein force field",        "Maier et al., 2015"),
    ("GAFF2",         "—",                  "Ligand force field",         "Wang et al., 2004"),
    ("ACPYPE",        "latest (pip)",       "Ligand parametrization",     "Sousa da Silva & Vranken, 2012"),
    ("AutoDock Vina", "1.2.6",              "Molecular docking",          "Eberhardt et al., 2021"),
    ("Meeko",         "0.7.1",              "Docking input preparation",  "Forli lab"),
    ("PROPKA3",       "3.5.1",              "pKa prediction",             "Olsson et al., 2011"),
    ("OpenBabel",     "3.x",               "File format conversion",      "O'Boyle et al., 2011"),
    ("RDKit",         "2025.x",             "Cheminformatics",            "Landrum et al., 2023"),
    ("Python",        "3.x",               "Scripting / analysis",       "Van Rossum & Drake, 2009"),
]

for row_data in rows_data:
    row = table.add_row()
    for cell, text in zip(row.cells, row_data):
        cell.text = text
        run = cell.paragraphs[0].runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)

doc.add_paragraph()

# ── References ──────────────────────────────────────────────────────
add_heading(doc, "References (Methods Section)", level=1)

refs = [
    "Abraham, M. J., Murtola, T., Schulz, R., Páll, S., Smith, J. C., Hess, B., & Lindahl, E. (2015). GROMACS: High performance molecular simulations through multi-level parallelism from laptops to supercomputers. SoftwareX, 1–2, 19–25. https://doi.org/10.1016/j.softx.2015.06.001",
    "Bussi, G., Donadio, D., & Parrinello, M. (2007). Canonical sampling through velocity rescaling. The Journal of Chemical Physics, 126(1), 014101. https://doi.org/10.1063/1.2408420",
    "Cho, Y., Gorina, S., Jeffrey, P. D., & Pavletich, N. P. (1994). Crystal structure of a p53 tumor suppressor-DNA complex: Understanding tumorigenic mutations. Science, 265(5170), 346–355. https://doi.org/10.1126/science.8023157",
    "Darden, T., York, D., & Pedersen, L. (1993). Particle mesh Ewald: An N·log(N) method for Ewald sums in large systems. The Journal of Chemical Physics, 98(12), 10089–10092. https://doi.org/10.1063/1.464397",
    "Eberhardt, J., Santos-Martins, D., Tillack, A. F., & Forli, S. (2021). AutoDock Vina 1.2.0: New docking methods, expanded force field, and Python bindings. Journal of Chemical Information and Modeling, 61(8), 3891–3898. https://doi.org/10.1021/acs.jcim.1c00203",
    "Essmann, U., Perera, L., Berkowitz, M. L., Darden, T., Lee, H., & Pedersen, L. G. (1995). A smooth particle mesh Ewald method. The Journal of Chemical Physics, 103(19), 8577–8593. https://doi.org/10.1063/1.470117",
    "Gowers, R. J., et al. (2016). MDAnalysis: A Python package for the rapid analysis of molecular dynamics simulations. Proceedings of the 15th Python in Science Conference (SciPy 2016), 98–105. https://doi.org/10.25080/Majora-629e541a-00e",
    "Hess, B., Bekker, H., Berendsen, H. J. C., & Fraaije, J. G. E. M. (1997). LINCS: A linear constraint solver for molecular simulations. Journal of Computational Chemistry, 18(12), 1463–1472. https://doi.org/10.1002/(SICI)1096-987X(199709)18:12<1463::AID-JCC4>3.0.CO;2-H",
    "Jakalian, A., Jack, D. B., & Bayly, C. I. (2002). Fast, efficient generation of high-quality atomic charges. AM1-BCC model: II. Parameterization and validation. Journal of Computational Chemistry, 23(16), 1623–1641. https://doi.org/10.1002/jcc.10128",
    "Jorgensen, W. L., Chandrasekhar, J., Madura, J. D., Impey, R. W., & Klein, M. L. (1983). Comparison of simple potential functions for simulating liquid water. The Journal of Chemical Physics, 79(2), 926–935. https://doi.org/10.1063/1.445869",
    "Maier, J. A., Martinez, C., Kasavajhala, K., Wickstrom, L., Hauser, K. E., & Simmerling, C. (2015). ff14SB: Improving the accuracy of protein side chain and backbone parameters from ff99SB. Journal of Chemical Theory and Computation, 11(8), 3696–3713. https://doi.org/10.1021/acs.jctc.5b00255",
    "Michaud-Agrawal, N., Denning, E. J., Woolf, T. B., & Beckstein, O. (2011). MDAnalysis: A toolkit for the analysis of molecular dynamics simulations. Journal of Computational Chemistry, 32(10), 2319–2327. https://doi.org/10.1002/jcc.21787",
    "O'Boyle, N. M., Banck, M., James, C. A., Morley, C., Vandermeersch, T., & Hutchison, G. R. (2011). Open Babel: An open chemical toolbox. Journal of Cheminformatics, 3, 33. https://doi.org/10.1186/1758-2946-3-33",
    "Olsson, M. H. M., Søndergaard, C. R., Rostkowski, M., & Jensen, J. H. (2011). PROPKA3: Consistent treatment of internal and surface residues in empirical pKa predictions. Journal of Chemical Theory and Computation, 7(2), 525–537. https://doi.org/10.1021/ct100578z",
    "Parrinello, M., & Rahman, A. (1981). Polymorphic transitions in single crystals: A new molecular dynamics method. Journal of Applied Physics, 52(12), 7182–7190. https://doi.org/10.1063/1.328693",
    "Shirts, M. R., Klein, C., Swails, J. M., Yin, J., Gilson, M. K., Mobley, D. L., Case, D. A., & Zhong, E. D. (2017). Lessons learned from comparing molecular dynamics engines on the SAMPL5 dataset. Journal of Computer-Aided Molecular Design, 31(1), 147–161. https://doi.org/10.1007/s10822-016-9977-1",
    "Sousa da Silva, A. W., & Vranken, W. F. (2012). ACPYPE - AnteChamber PYthon Parser interfacE. BMC Research Notes, 5, 367. https://doi.org/10.1186/1756-0500-5-367",
    "Wang, J., Wolf, R. M., Caldwell, J. W., Kollman, P. A., & Case, D. A. (2004). Development and testing of a general amber force field. Journal of Computational Chemistry, 25(9), 1157–1174. https://doi.org/10.1002/jcc.20035",
]

for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)  # hanging indent
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)

# ── Save ─────────────────────────────────────────────────────────────
output_path = "/Users/erdalbalcan/md_1tsr_carvacrol/Methods_1TSR_Carvacrol_MD.docx"
doc.save(output_path)
print(f"Word dosyası kaydedildi: {output_path}")
